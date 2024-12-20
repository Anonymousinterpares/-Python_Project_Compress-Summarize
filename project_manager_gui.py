import sys
import os
from pathlib import Path
from datetime import datetime
import json
import logging
from PyQt5.QtWidgets import (
    QApplication,
    QWidget,
    QPushButton,
    QFileDialog,
    QMessageBox,
    QMainWindow,
    QAction,
    QInputDialog,
    QCheckBox,
    QGroupBox,
    QVBoxLayout,
    QHBoxLayout,
    QRadioButton,
    QLabel,
    QLineEdit,
    QDialog,
    QSlider,
    QComboBox,
)
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import openai  # Import the openai library
import google.generativeai as genai # Import the google generativeai library

def setup_logging():
    script_dir = os.path.dirname(os.path.abspath(__file__))  # Get the directory of the script
    log_dir = os.path.join(script_dir, "logs")  # Create a logs subdirectory
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    log_file = os.path.join(log_dir, "dev_manager.log")

    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)

    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.DEBUG)

    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    file_handler.setFormatter(formatter)

    logger.addHandler(file_handler)

    # Add a StreamHandler to output logs to the console as well
    stream_handler = logging.StreamHandler()
    stream_handler.setLevel(logging.INFO)  # Set the level for console output
    stream_handler.setFormatter(formatter)
    logger.addHandler(stream_handler)

    logging.info(f"Logging setup complete. Log file: {log_file}")
    print(f"Logging setup complete. Log file: {log_file}")

setup_logging()

class APISettingsDialog(QDialog):
    """Dialog for editing API settings."""
    def __init__(self, settings, parent=None):
        super().__init__(parent)
        self.setWindowTitle('API Settings')
        self.settings = settings

        self.providers = ["OpenAI", "Google"]
        self.provider_combobox = QComboBox()
        self.provider_combobox.addItems(self.providers)
        self.provider_combobox.currentTextChanged.connect(self.update_api_fields)

        self.api_key_label = QLabel("API Key:")
        self.api_key_edit = QLineEdit()
        self.api_key_edit.setEchoMode(QLineEdit.Password)

        self.temperature_label = QLabel("Temperature:")
        self.temperature_slider = QSlider(Qt.Horizontal)
        self.temperature_slider.setMinimum(0)
        self.temperature_slider.setMaximum(100)
        self.temperature_slider.setValue(int(self.settings.get("temperature", 70) * 100)) # Default to 0.7
        self.temperature_display_label = QLabel(f"{self.temperature_slider.value() / 100:.2f}")
        self.temperature_slider.valueChanged.connect(self.update_temperature_display)

        self.model_label = QLabel("Model:")
        self.model_combobox = QComboBox()

        self.load_button = QPushButton("Load Settings")
        self.load_button.clicked.connect(self.load_settings_for_provider)
        self.save_button = QPushButton("Save Settings")
        self.save_button.clicked.connect(self.save_settings)
        self.reset_button = QPushButton("Reset Settings")
        self.reset_button.clicked.connect(self.reset_settings)

        layout = QVBoxLayout()
        layout.addWidget(self.provider_combobox)
        layout.addWidget(self.api_key_label)
        layout.addWidget(self.api_key_edit)

        temp_layout = QHBoxLayout()
        temp_layout.addWidget(self.temperature_label)
        temp_layout.addWidget(self.temperature_slider)
        temp_layout.addWidget(self.temperature_display_label)
        layout.addLayout(temp_layout)

        layout.addWidget(self.model_label)
        layout.addWidget(self.model_combobox)

        button_layout = QHBoxLayout()
        button_layout.addWidget(self.load_button)
        button_layout.addWidget(self.save_button)
        button_layout.addWidget(self.reset_button)
        layout.addLayout(button_layout)

        self.setLayout(layout)
        self.update_api_fields(self.provider_combobox.currentText())
        self.load_settings_for_provider(self.provider_combobox.currentText())

    def update_temperature_display(self, value):
        self.temperature_display_label.setText(f"{value / 100:.2f}")

    def update_api_fields(self, provider):
        self.api_key_label.setText(f"{provider} API Key:")
        self.api_key_edit.clear()
        self.model_combobox.clear()
        if provider == "OpenAI":
            self.model_combobox.addItems(["gpt-4o", "gpt-4o-mini", "gpt-4", "gpt-3.5-turbo"])
        elif provider == "Google":
            self.model_combobox.addItems(["gemini-pro"]) # Add actual Google models

    def load_settings_for_provider(self, provider):
        provider_settings = self.settings.get(provider.lower(), {})
        api_key = provider_settings.get("api_key", "")
        if api_key:
            self.api_key_edit.setText(api_key[:10] + "*" * (len(api_key) - 10))
        self.temperature_slider.setValue(int(provider_settings.get("temperature", 0.7) * 100))
        self.model_combobox.setCurrentText(provider_settings.get("model", ""))

    def save_settings(self):
        """Save the API settings."""
        provider = self.provider_combobox.currentText()
        api_key = self.api_key_edit.text()
        temperature = self.temperature_slider.value() / 100
        model = self.model_combobox.currentText()

        if provider not in self.settings:
            self.settings[provider.lower()] = {}

        self.settings[provider.lower()]["api_key"] = api_key
        self.settings[provider.lower()]["temperature"] = temperature
        self.settings[provider.lower()]["model"] = model

        try:
            # Save settings.json in the same directory as the script
            script_dir = os.path.dirname(os.path.abspath(__file__))
            settings_path = os.path.join(script_dir, "settings.json")
            with open(settings_path, "w") as f:
                json.dump(self.settings, f, indent=4)
            QMessageBox.information(self, "Settings Saved", "API settings saved successfully.")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error saving settings: {e}")

    def reset_settings(self):
        """Reset the API settings for the current provider."""
        provider = self.provider_combobox.currentText()
        reply = QMessageBox.warning(
            self, "Warning",
            f"Are you sure you want to reset the API settings for {provider}?\nThis action is irreversible.",
            QMessageBox.Yes | QMessageBox.Cancel
        )
        if reply == QMessageBox.Yes:
            if provider.lower() in self.settings:
                del self.settings[provider.lower()]
                try:
                    script_dir = os.path.dirname(os.path.abspath(__file__))
                    settings_path = os.path.join(script_dir, "settings.json")
                    with open(settings_path, "w") as f:
                        json.dump(self.settings, f, indent=4)
                    QMessageBox.information(self, "Settings Reset", f"{provider} API settings have been reset.")
                    self.load_settings_for_provider(provider) # Clear the fields
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error resetting settings: {e}")
            else:
                QMessageBox.information(self, "No Settings", f"No settings found for {provider}.")

class ProjectManagerGUI(QMainWindow):
    """Main application window."""
    def __init__(self):
        super().__init__()
        self.api_settings = self.load_api_settings()
        self.project_path = None  # Store the selected project path
        self.doc_file_path = None  # Store the selected documentation file path
        self.initUI()

    def initUI(self):
        # ... (Previous GUI setup) ...
        """Initialize the UI elements."""
        self.setWindowTitle('Project Manager')
        self.setGeometry(100, 100, 600, 400)

        # Central Widget
        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout()
        central_widget.setLayout(layout)

        # Project/Documentation Selection
        self.select_project_button = QPushButton('Select Project Folder', self)
        self.select_project_button.clicked.connect(self.select_folder)
        layout.addWidget(self.select_project_button)

        self.select_doc_button = QPushButton('Select Documentation File', self)
        self.select_doc_button.clicked.connect(self.select_file)
        layout.addWidget(self.select_doc_button)

        # File Format Group
        format_group = QGroupBox("File Format")
        format_layout = QVBoxLayout()

        self.docx_radio = QRadioButton("DOCX Format")
        self.txt_radio = QRadioButton("TXT Format")
        self.txt_radio.setChecked(True)  # Default to TXT
        self.docx_radio.toggled.connect(self.update_action_state)
        self.txt_radio.toggled.connect(self.update_action_state)

        format_layout.addWidget(self.docx_radio)
        format_layout.addWidget(self.txt_radio)
        format_group.setLayout(format_layout)
        layout.addWidget(format_group)

        # Action Group
        action_group = QGroupBox("Action")
        action_layout = QVBoxLayout()

        self.compress_radio = QRadioButton("Compress Project")
        self.reconstruct_radio = QRadioButton("Reconstruct Project")
        self.compress_radio.setChecked(True)  # Default to Compress
        action_layout.addWidget(self.compress_radio)
        action_layout.addWidget(self.reconstruct_radio)
        action_group.setLayout(action_layout)
        layout.addWidget(action_group)

        # LLM Options
        llm_group = QGroupBox("LLM Options")
        llm_layout = QVBoxLayout()

        self.use_llm_check = QCheckBox("Use LLM for Documentation")
        llm_layout.addWidget(self.use_llm_check)

        self.llm_type_label = QLabel("LLM Type:")
        self.none_radio = QRadioButton("None")
        self.openai_radio = QRadioButton("OpenAI")
        self.google_radio = QRadioButton("Google")
        self.none_radio.setChecked(True)
        llm_layout.addWidget(self.llm_type_label)
        llm_layout.addWidget(self.none_radio)
        llm_layout.addWidget(self.openai_radio)
        llm_layout.addWidget(self.google_radio)

        self.overview_label = QLabel("Overview Type:")
        self.general_radio = QRadioButton("General")
        self.detailed_radio = QRadioButton("Detailed")
        self.general_radio.setChecked(True)
        llm_layout.addWidget(self.overview_label)
        llm_layout.addWidget(self.general_radio)
        llm_layout.addWidget(self.detailed_radio)

        llm_group.setLayout(llm_layout)
        layout.addWidget(llm_group)

        # Process Button
        self.process_button = QPushButton('Process', self)
        self.process_button.clicked.connect(self.process_project)
        layout.addWidget(self.process_button)

        # Menu Bar
        menubar = self.menuBar()
        file_menu = menubar.addMenu('&File')
        settings_menu = menubar.addMenu('&Settings')

        # Exit action
        exit_action = QAction('Exit', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # API Settings action
        api_settings_action = QAction('API Settings', self)
        api_settings_action.triggered.connect(self.show_api_settings)
        settings_menu.addAction(api_settings_action)

        self.update_action_state()
        self.show()

    def select_folder(self):
        """Handle project folder selection."""
        folder_path = QFileDialog.getExistingDirectory(self, "Select Project Folder")
        if folder_path:
            self.project_path = folder_path
            print(f"Selected folder: {self.project_path}")

    def select_file(self):
        """Handle documentation file selection."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Documentation File",
            "",
            "Documentation files (*.txt *.docx);;All files (*.*)",
        )
        if file_path:
            self.doc_file_path = file_path
            print(f"Selected file: {self.doc_file_path}")

    def update_action_state(self):
        """Updates the state of the action radio buttons based on the selected file format."""
        if self.docx_radio.isChecked():
            self.compress_radio.setChecked(True)
            self.compress_radio.setEnabled(False)
            self.reconstruct_radio.setEnabled(True)
        elif self.txt_radio.isChecked():
            self.compress_radio.setEnabled(True)
            self.reconstruct_radio.setEnabled(True)

    def process_project(self):
        """Handle the main project processing logic."""
        logging.info("Starting project processing")

        if self.compress_radio.isChecked():
            logging.info("Compress project option selected")
            # Handle compression (documentation generation)
            if not self.project_path:
                QMessageBox.warning(self, "Error", "Please select a project folder first.")
                logging.warning("No project folder selected for compression")
                return

            llm_output = None
            if self.use_llm_check.isChecked():
                logging.info("Fetching LLM documentation...")
                project_text_content_for_llm = self.get_project_content_for_llm(self.project_path)
                if project_text_content_for_llm:
                    llm_output = self.generate_llm_documentation(project_text_content_for_llm)
                    if llm_output:
                        logging.info("LLM documentation fetched successfully.")
                    else:
                        logging.warning("LLM documentation generation failed.")
                else:
                    logging.warning("Could not extract project content for LLM.")

            if self.txt_radio.isChecked():
                logging.info("TXT format selected")
                output_file = os.path.join(
                    self.project_path, "project_documentation.txt"
                ).replace("\\", "/")
                self.convert_project_to_text(self.project_path, output_file, llm_overview=llm_output)
            elif self.docx_radio.isChecked():
                logging.info("DOCX format selected")
                output_file = os.path.join(
                    self.project_path, "project_documentation.docx"
                ).replace("\\", "/")
                self.create_project_documentation(self.project_path, output_file, llm_content=llm_output)

        elif self.reconstruct_radio.isChecked():
            logging.info("Reconstruct project option selected")
            # Handle reconstruction
            if not self.doc_file_path:
                QMessageBox.warning(
                    self, "Error", "Please select a documentation file first."
                )
                logging.warning("No documentation file selected for reconstruction")
                return

            if self.doc_file_path.endswith(".txt"):
                logging.info("TXT format selected for reconstruction")
                project_name = QInputDialog.getText(
                    self, "Project Name", "Enter the name for the reconstructed project:"
                )[0]
                if project_name:
                    save_location = QFileDialog.getExistingDirectory(
                        self, "Select Save Location"
                    )
                    if save_location:
                        # Use forward slashes for consistency
                        save_location = save_location.replace("\\", "/")
                        self.recreate_project_from_text(
                            self.doc_file_path, project_name, save_location
                        )
            else:
                QMessageBox.warning(
                    self,
                    "Error",
                    "Invalid documentation file type for reconstruction. Select a .txt file.",
                )
                logging.warning(
                    f"Invalid documentation file type selected for reconstruction: {self.doc_file_path}"
                )
    def get_project_content_for_llm(self, project_path):
        """Extracts content from project files for LLM processing."""
        project_content_for_llm = ""
        for root, _, files in os.walk(project_path):
            for file in files:
                if file.endswith((".py", ".json", ".log")) and not file.startswith("."):
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, project_path)
                    try:
                        with open(file_path, "r", encoding="utf-8") as sourcefile:
                            content = sourcefile.read()
                            project_content_for_llm += f"File: {rel_path}\n{content}\n\n"
                    except Exception as e:
                        logging.error(f"Error reading file for LLM: {file_path} - {e}")
        return project_content_for_llm
    
    def generate_llm_documentation(self, project_text):
        """Generates documentation using the selected LLM."""
        selected_llm = None
        if self.openai_radio.isChecked():
            selected_llm = "openai"
        elif self.google_radio.isChecked():
            selected_llm = "google"

        if not selected_llm:
            QMessageBox.warning(self, "Warning", "Please select an LLM provider.")
            return None

        llm_settings = self.api_settings.get(selected_llm, {})
        api_key = llm_settings.get("api_key")
        model = llm_settings.get("model")
        temperature = llm_settings.get("temperature", 0.7)

        if not api_key:
            QMessageBox.warning(self, "Error", f"{selected_llm.capitalize()} API key not configured.")
            return None

        overview_type = "general" if self.general_radio.isChecked() else "detailed"

        system_prompt = f"You are a Coding Master tasked with explaining the dependencies of the following code in a clear and structured way. Provide a {overview_type} overview."

        if selected_llm == "openai":
            return self.call_openai_api(api_key, model, system_prompt, project_text, temperature)
        elif selected_llm == "google":
            return self.call_google_api(api_key, model, system_prompt, project_text, temperature) # Added model parameter for consistency

    def call_openai_api(self, api_key, model, system_prompt, content, temperature):
        """Calls the OpenAI API to generate documentation."""
        logging.info(f"Calling OpenAI API with model: {model}, temperature: {temperature}")
        openai.api_key = api_key
        try:
            logging.debug(f"OpenAI API Request - Model: {model}, Temperature: {temperature}, Prompt: {system_prompt}, Content: {content[:500]}...") # Log first 500 chars of content
            response = openai.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": content}
                ],
                temperature=temperature,
            )
            logging.debug(f"OpenAI API Response: {response}")
            return response.choices[0].message.content
        except Exception as e:
            logging.error(f"Error calling OpenAI API: {e}")
            QMessageBox.critical(self, "OpenAI Error", f"Error communicating with OpenAI: {e}")
            return None

    def call_google_api(self, api_key, model, system_prompt, content, temperature):
        """Calls the Google API to generate documentation."""
        logging.info(f"Calling Google API with model: {model}, temperature: {temperature}")
        genai.configure(api_key=api_key)
        generation_config = genai.types.GenerationConfig(
            temperature=temperature
        )
        model = genai.GenerativeModel(model_name=model,
                                     generation_config=generation_config)

        prompt = f"{system_prompt}\n\n{content}"
        logging.debug(f"Google API Request - Model: {model.name}, Temperature: {temperature}, Prompt: {prompt[:500]}...") # Log first 500 chars of prompt
        try:
            response = model.generate_content(prompt)
            logging.debug(f"Google API Response: {response.text}")
            return response.text
        except Exception as e:
            logging.error(f"Error calling Google API: {e}")
            QMessageBox.critical(self, "Google AI Error", f"Error communicating with Google AI: {e}")
            return None

    def show_api_settings(self):
        """Display the API settings dialog."""
        dialog = APISettingsDialog(self.api_settings, self)
        if dialog.exec_() == QDialog.Accepted:
            # Reload settings if saved
            self.api_settings = self.load_api_settings()

    def load_api_settings(self):
        """Load API settings from settings.json."""
        try:
            script_dir = os.path.dirname(os.path.abspath(__file__))
            settings_path = os.path.join(script_dir, "settings.json")
            with open(settings_path, "r") as f:
                return json.load(f)
        except FileNotFoundError:
            return {}

    def convert_project_to_text(self, project_path, output_file, llm_overview=None):
        """
        Converts all Python, JSON, and LOG files in a project to a single text file,
        including directory structure and handling incompatible files.

        Args:
            project_path (str): Path to the project root directory.
            output_file (str):  Intended name of the output file (will be modified to include timestamp).
            llm_overview (str, optional): LLM-generated overview of the project. Defaults to None.
        """
        if not project_path:
            QMessageBox.warning(self, "Error", "No project folder selected.")
            return None  # Stop execution if no folder is selected

        timestamp = datetime.now().strftime("%d.%m.%Y_%H_%M_%S")
        output_file = os.path.join(
            project_path, f"project_documentation_{timestamp}.txt"
        )

        def get_project_structure(directory):
            """Returns a string representation of the project structure."""
            structure = []
            base_path = os.path.basename(directory)
            structure.append(f"{base_path}/")

            for root, dirs, files in os.walk(directory):
                dirs[:] = [
                    d for d in dirs if d != "__pycache__" and not d.startswith(".")
                ]

                rel_path = os.path.relpath(root, directory)
                if rel_path == ".":
                    continue

                level = rel_path.count(os.sep) + 1
                indent = "  " * level
                folder = os.path.basename(root)
                structure.append(f"{indent}{folder}/")

                sub_indent = "  " * (level + 1)
                for file in sorted(files):
                    if not file.startswith("."):
                        structure.append(f"{sub_indent}{file}")

                return "\n".join(structure)

        incompatible_files = []
        incompatible_structure = []

        try:
            with open(output_file, "w", encoding="utf-8") as outfile:
                outfile.write(f"Project Documentation\n")
                outfile.write(
                    f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                )

                if llm_overview:
                    outfile.write("\n" + "=" * 50 + "\n")
                    outfile.write("PROJECT GENERAL OVERVIEW:\n")
                    outfile.write("=" * 50 + "\n\n")
                    outfile.write(llm_overview + "\n\n")

                outfile.write("=" * 50 + "\n\n")
                outfile.write("PROJECT STRUCTURE:\n")
                outfile.write("=" * 20 + "\n")
                structure = get_project_structure(project_path)
                outfile.write(structure)
                outfile.write("\n\n" + "=" * 50 + "\n\n")

                outfile.write("FILES CONTENT:\n")
                outfile.write("=" * 20 + "\n\n")

                for root, dirs, files in os.walk(project_path):
                    dirs[:] = [
                        d for d in dirs if d != "__pycache__" and not d.startswith(".")
                    ]
                    for file in files:
                        if file.startswith("."):
                            continue

                        file_path = os.path.join(root, file)
                        rel_path = os.path.relpath(file_path, project_path)

                        if file.endswith((".py", ".json", ".log")):
                            outfile.write(f"File: {rel_path}\n")
                            outfile.write("-" * len(f"File: {rel_path}") + "\n\n")

                            try:
                                with open(file_path, "r", encoding="utf-8") as sourcefile:
                                    content = sourcefile.read()
                                    outfile.write(content)
                                    outfile.write("\n\n" + "=" * 50 + "\n\n")
                            except Exception as e:
                                outfile.write(f"Error reading file: {str(e)}\n\n")
                        else:
                            incompatible_files.append(rel_path)
                            level = rel_path.count(os.sep)
                            indent = "  " * level
                            incompatible_structure.append(
                                f"{indent}{os.path.basename(rel_path)}"
                            )

            if incompatible_files:
                incompatible_file_path = os.path.join(
                    project_path, "incompatible_files.txt"
                )
                with open(
                    incompatible_file_path, "w", encoding="utf-8"
                ) as incompatible_file:
                    incompatible_file.write("INCOMPATIBLE FILES STRUCTURE:\n")
                    incompatible_file.write("=" * 30 + "\n")
                    incompatible_file.write("\n".join(incompatible_structure))
                    incompatible_file.write("\n\n")
                    incompatible_file.write("Incompatible files:\n")
                    incompatible_file.write("\n".join(incompatible_files))

                msg = QMessageBox()
                msg.setIcon(QMessageBox.Information)
                msg.setText(
                    "Documentation generated with incompatible files."
                )
                msg.setInformativeText(f"Details in: {incompatible_file_path}")
                msg.setWindowTitle("Info")
                msg.exec_()

            else:
                msg = QMessageBox()
                msg.setIcon(QMessageBox.Information)
                msg.setText("Documentation generated successfully!")
                msg.setInformativeText(f"Saved to: {output_file}")
                msg.setWindowTitle("Success")
                msg.exec_()
            return True  # Indicate success

        except Exception as e:
            msg = QMessageBox()
            msg.setIcon(QMessageBox.Critical)
            msg.setText("An error occurred!")
            msg.setInformativeText(f"{str(e)}")
            msg.setWindowTitle("Error")
            msg.exec_()
            return False # Indicate failure

# Modified to include a timestamp in the output filename.
# The filename now includes the current date and time in DD.MM.YYYY_HH_MM_SS format.

    def recreate_project_from_text(self, doc_file, project_name, save_location):
            """
            Recreates a project structure and files from a documentation text file.

            Args:
                doc_file (str): Path to the documentation file.
                project_name (str): Name of the project folder to create.
                save_location (str): Location where to create the project.
            """
            logging.info(
                f"Recreating project from TXT: {doc_file}, Project Name: {project_name}, Save Location: {save_location}"
            )
            # Use forward slashes for consistency
            project_path = os.path.join(save_location, project_name).replace("\\", "/")

            try:
                with open(doc_file, "r", encoding="utf-8") as f:
                    content = f.read()

                structure_start = (
                    content.find("PROJECT STRUCTURE:\n")
                    + len("PROJECT STRUCTURE:\n")
                    + len("=" * 20 + "\n")
                )
                structure_end = content.find("\n\n" + "=" * 50, structure_start)
                structure_section = content[structure_start:structure_end].strip()

                current_dir_stack = []
                for line in structure_section.split("\n"):
                    stripped_line = line.strip()
                    indent_level = (len(line) - len(stripped_line)) // 2
                    logging.info(
                        f"Processing line: '{line}', Stripped: '{stripped_line}', Indent Level: {indent_level}"
                    )

                    while len(current_dir_stack) > indent_level:
                        current_dir_stack.pop()

                    if stripped_line.endswith("/"):
                        dir_name = stripped_line.rstrip("/")
                        current_dir_stack.append(dir_name)
                        if indent_level == 0:
                            continue
                        # Use forward slashes for consistency
                        dir_path = os.path.join(project_path, *current_dir_stack).replace("\\", "/")
                        logging.info(f"Creating directory: {dir_path}")
                        os.makedirs(dir_path, exist_ok=True)
                    else:
                        if current_dir_stack:
                            # Use forward slashes for consistency
                            file_path = os.path.join(
                                project_path, *current_dir_stack, stripped_line
                            ).replace("\\", "/")
                        else:
                            # Use forward slashes for consistency
                            file_path = os.path.join(project_path, stripped_line).replace("\\", "/")
                        logging.info(f"Creating file: {file_path}")
                        os.makedirs(os.path.dirname(file_path), exist_ok=True)
                        Path(file_path).touch()

                files_section = content[content.find("FILES CONTENT:\n") :]
                current_file = None
                current_content = []
                in_file_content = False

                for line in files_section.split("\n"):
                    if line.startswith("File: "):
                        if current_file and current_content:
                            # Use forward slashes for consistency
                            file_path = os.path.join(project_path, current_file).replace("\\", "/")
                            logging.info(f"Writing content to file: {file_path}")
                            os.makedirs(os.path.dirname(file_path), exist_ok=True)

                            content_str = "\n".join(current_content).strip()
                            with open(file_path, "w", encoding="utf-8") as f:
                                if file_path.endswith(".json"):
                                    try:
                                        json_content = json.loads(content_str)
                                        json.dump(json_content, f, indent=4)
                                    except json.JSONDecodeError:
                                        f.write(content_str)
                                else:
                                    f.write(content_str)

                        current_file = line[6:].strip()
                        current_content = []
                        in_file_content = True
                    elif line.startswith("=" * 50):
                        in_file_content = False
                    elif line.startswith("-" * len("File: ")):
                        continue
                    elif in_file_content:
                        current_content.append(line)

                if current_file and current_content:
                    # Use forward slashes for consistency
                    file_path = os.path.join(project_path, current_file).replace("\\", "/")
                    logging.info(f"Writing content to last file: {file_path}")
                    os.makedirs(os.path.dirname(file_path), exist_ok=True)

                    content_str = "\n".join(current_content).strip()
                    with open(file_path, "w", encoding="utf-8") as f:
                        if file_path.endswith(".json"):
                            try:
                                json_content = json.loads(content_str)
                                json.dump(json_content, f, indent=4)
                            except json.JSONDecodeError:
                                f.write(content_str)
                        else:
                            f.write(content_str)

                # --- Correct QMessageBox usage ---
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Information)
                msg.setText("Project recreated successfully!")
                msg.setInformativeText(f"Location: {project_path}")
                msg.setWindowTitle("Success")
                msg.exec_()
                logging.info(f"Project recreated successfully at: {project_path}")

            except Exception as e:
                # --- Correct QMessageBox usage ---
                msg = QMessageBox(self)
                msg.setIcon(QMessageBox.Critical)
                msg.setText(f"An error occurred: {str(e)}")
                msg.setWindowTitle("Error")
                msg.exec_()
                logging.error(f"Error during project recreation: {str(e)}")
    def create_project_documentation(self, project_path, output_file, llm_content=None):
        """
        Creates detailed project documentation in DOCX format with structured sections
        and handles incompatible files.
        Args:
            project_path (str): Path to the project root directory.
            output_file (str): Intended name of the output file (will be modified to include timestamp).
            llm_content (str, optional): LLM-generated content to include in the documentation. Defaults to None.
        """
        logging.info(
            f"Creating DOCX project documentation for: {project_path}"
        )
        timestamp = datetime.now().strftime("%d.%m.%Y_%H_%M_%S")
        output_file = os.path.join(
            project_path, f"project_documentation_{timestamp}.docx"
        )
        logging.info(f"Output file: {output_file}")
        doc = Document()

        # Define styles
        styles = doc.styles

        style_normal = styles["Normal"]
        style_normal.font.name = "Calibri"
        style_normal.font.size = Pt(11)

        for level in range(1, 4):
            style_name = f"Custom Heading {level}"
            if style_name not in styles:
                style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = styles[f"Heading {level}"]
                style.font.name = "Calibri"
                style.font.size = Pt(16 - level)
                style.font.bold = True

        doc.add_heading("Project Documentation", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Project Path: {project_path}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("Document sections:", style="List Bullet")
        sections = [
            "1. Project Overview",
            "2. Project Structure",
            "3. Code and Log Files Documentation",
            "4. Incompatible Files",
            "5. Dependencies",
            "6. Setup Instructions",
        ]
        if llm_content:
            sections.insert(4, "7. LLM Analysis") # Insert before Dependencies
        for section in sections:
            doc.add_paragraph(section, style="List Number")
        doc.add_page_break()

        doc.add_heading("1. Project Overview", level=1)
        doc.add_paragraph(
            "This documentation provides a comprehensive overview of the project structure and contents."
        )
        doc.add_paragraph("Project Details:", style="Custom Heading 3")
        project_name = os.path.basename(project_path)
        details = [
            f"Project Name: {project_name}",
            f'Documentation Date: {datetime.now().strftime("%Y-%m-%d")}',
            f'Number of Python Files: {sum(1 for root, _, files in os.walk(project_path) for file in files if file.endswith(".py"))}',
            f'Number of Log Files: {sum(1 for root, _, files in os.walk(project_path) for file in files if file.endswith(".log"))}',
        ]
        for detail in details:
            doc.add_paragraph(detail, style="List Bullet")

        doc.add_heading("2. Project Structure", level=1)
        doc.add_paragraph("Directory structure of the project:", style="Custom Heading 3")

        structure = []
        incompatible_files = []
        incompatible_structure = []
        for root, dirs, files in os.walk(project_path):
            dirs[:] = [d for d in dirs if d != "__pycache__" and not d.startswith(".")]
            level = root.replace(project_path, "").count(os.sep)
            indent = "    " * level
            folder = os.path.basename(root)
            structure.append(f"{indent}{folder}/")
            for file in sorted(files):
                if not file.startswith("."):
                    if file.endswith((".py", ".json", ".log")):
                        structure.append(f"{indent}    {file}")
                    else:
                        incompatible_files.append(os.path.relpath(os.path.join(root, file), project_path))
                        incompatible_structure.append(f"{indent}    {file}")
        doc.add_paragraph().add_run("\n".join(structure)).font.name = "Courier New"

        doc.add_heading("3. Code and Log Files Documentation", level=1)

        for root, _, files in os.walk(project_path):
            for file in files:
                if file.endswith((".py", ".log")) and not file.startswith("."):
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, project_path)

                    doc.add_heading(f"File: {rel_path}", level=2)

                    try:
                        with open(file_path, "r", encoding="utf-8") as pyfile:
                            content = pyfile.read()
                            p = doc.add_paragraph()
                            run = p.add_run(content)
                            run.font.name = "Courier New"
                            run.font.size = Pt(9)
                    except Exception as e:
                        doc.add_paragraph(f"Error reading file: {str(e)}")

                    doc.add_paragraph()

        doc.add_heading("4. Incompatible Files", level=1)
        if incompatible_files:
            doc.add_paragraph("List of incompatible files:", style="Custom Heading 3")
            doc.add_paragraph().add_run("\n".join(incompatible_structure)).font.name = "Courier New"
            doc.add_paragraph("Incompatible files details:", style="List Bullet")
            for file_path in incompatible_files:
                doc.add_paragraph(file_path, style="List Bullet")
        else:
            doc.add_paragraph("No incompatible files found.")

        if llm_content:
            doc.add_heading("7. LLM Analysis", level=1)
            doc.add_paragraph(llm_content)
        else:
            doc.add_paragraph("LLM documentation was not requested for this document.")

        doc.add_heading("5. Dependencies", level=1)
        doc.add_paragraph("List of potential project dependencies:", style="Custom Heading 3")
        doc.add_paragraph("To be filled manually with:")
        dependencies = [
            "Required Python version",
            "Required external packages",
            "System requirements",
            "Additional software dependencies",
        ]
        for dep in dependencies:
            doc.add_paragraph(dep, style="List Bullet")

        doc.add_heading("6. Setup Instructions", level=1)
        doc.add_paragraph("Template for setup instructions:", style="Custom Heading 3")
        instructions = [
            "Environment setup",
            "Installation steps",
            "Configuration requirements",
            "Running the project",
            "Testing procedures",
        ]
        for instruction in instructions:
            doc.add_paragraph(instruction, style="List Bullet")

        try:
            doc.save(output_file)
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Information)
            msg.setText("Documentation generated successfully!")
            msg.setInformativeText(f"Saved to: {output_file}")
            msg.setWindowTitle("Success")
            msg.exec_()
            logging.info(f"DOCX documentation generated successfully at: {output_file}")
        except Exception as e:
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Critical)
            msg.setText(f"An error occurred while saving: {str(e)}")
            msg.setWindowTitle("Error")
            msg.exec_()
            logging.error(f"Error during DOCX generation: {str(e)}")

    def recreate_project_from_docx(self, doc_file, project_name, save_location):
        """
        Recreates a project from a DOCX file by first converting it to TXT.

        Args:
            doc_file (str): Path to the DOCX documentation file.
            project_name (str): Name of the project folder to create.
            save_location (str): Location where to create the project.
        """
        logging.info(
            f"Recreating project from DOCX: {doc_file}, Project Name: {project_name}, Save Location: {save_location}"
        )

        try:
            # Convert DOCX to TXT
            txt_file = self.convert_docx_to_txt(doc_file)
            if txt_file is None:
                logging.error("Failed to convert DOCX to TXT.")
                return
            logging.info(f"DOCX file converted to TXT: {txt_file}")

            # Recreate project from TXT
            self.recreate_project_from_text(txt_file, project_name, save_location)
            logging.info(f"Project recreated from TXT: {txt_file}")

        except Exception as e:
            logging.error(f"Error during project recreation from DOCX: {str(e)}")
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Critical)
            msg.setText(f"An error occurred: {str(e)}")
            msg.setWindowTitle("Error")
            msg.exec_()

    def convert_docx_to_txt(self, docx_path):
        """
        Converts a DOCX file to a TXT file.

        Args:
            docx_path (str): Path to the DOCX file.

        Returns:
            str: Path to the converted TXT file, or None if an error occurred.
        """
        try:
            doc = Document(docx_path)
            txt_path = os.path.join(
                os.path.dirname(docx_path),
                os.path.splitext(os.path.basename(docx_path))[0] + ".txt"
            )

            with open(txt_path, "w", encoding="utf-8") as txt_file:
                for paragraph in doc.paragraphs:
                    txt_file.write(paragraph.text + "\n")

            return txt_path

        except Exception as e:
            logging.error(f"Error converting DOCX to TXT: {str(e)}")
            return None

    def call_openai_api(self, api_key, model, system_prompt, content, temperature):
        """Calls the OpenAI API to generate documentation."""
        logging.info(f"Calling OpenAI API with model: {model}, temperature: {temperature}")
        openai.api_key = api_key
        try:
            logging.debug(f"OpenAI API Request - Model: {model}, Temperature: {temperature}, Prompt: {system_prompt}, Content: {content[:500]}...") # Log first 500 chars of content
            response = openai.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": content}
                ],
                temperature=temperature,
            )
            logging.debug(f"OpenAI API Response: {response}")
            return response.choices[0].message.content
        except Exception as e:
            logging.error(f"Error calling OpenAI API: {e}")
            QMessageBox.critical(self, "OpenAI Error", f"Error communicating with OpenAI: {e}")
            return None

    def call_google_api(self, api_key, model, system_prompt, content, temperature):
        """Calls the Google API to generate documentation."""
        logging.info(f"Calling Google API with model: {model}, temperature: {temperature}")
        genai.configure(api_key=api_key)
        generation_config = genai.types.GenerationConfig(
            temperature=temperature
        )
        model = genai.GenerativeModel(model_name=model,
                                     generation_config=generation_config)

        prompt = f"{system_prompt}\n\n{content}"
        logging.debug(f"Google API Request - Model: {model.name}, Temperature: {temperature}, Prompt: {prompt[:500]}...") # Log first 500 chars of prompt
        try:
            response = model.generate_content(prompt)
            logging.debug(f"Google API Response: {response.text}")
            return response.text
        except Exception as e:
            logging.error(f"Error calling Google API: {e}")
            QMessageBox.critical(self, "Google AI Error", f"Error communicating with Google AI: {e}")
            return None

if __name__ == "__main__":
    app = QApplication(sys.argv)
    ex = ProjectManagerGUI()
    sys.exit(app.exec_())
