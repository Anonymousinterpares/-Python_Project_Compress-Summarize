import os
from PyQt5.QtWidgets import QFileDialog, QMessageBox, QInputDialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

from core.utils import has_any_extension, _read_file_content  # Import _read_file_content
from utils.file_utils import get_base_dir

def select_file(main_window):
    """Handle documentation file selection, remembering the last location."""
    file_path, _ = QFileDialog.getOpenFileName(
        main_window,
        "Select Documentation File",
        main_window.recent_doc_path,
        "Documentation files (*.txt *.docx);;All files (*.*)",
    )
    if file_path:
        main_window.doc_file_path = file_path
        main_window.recent_doc_path = os.path.dirname(file_path)
        print(f"Selected file: {main_window.doc_file_path}")

def reset_doc_selection(main_window):
    """Resets the documentation file selection."""
    main_window.doc_file_path = None
    main_window.recent_doc_path = ""
    print("Documentation file selection reset.")
    QMessageBox.information(main_window, "Reset", "Documentation file selection has been reset.")

def convert_project_to_text(main_window, project_path, llm_overview=None):
    """
    Converts project files to a single text documentation file in Markdown format,
    handling various extensions and recognizing code blocks.
    """

    def get_project_structure(directory):
        """Creates a string representation of the project directory structure."""
        structure = []
        base_path = os.path.basename(directory)
        structure.append(f"{base_path}/")

        for root, dirs, files in os.walk(directory):
            dirs[:] = [d for d in dirs if d != "__pycache__" and not d.startswith(".")]
            level = os.path.relpath(root, directory).count(os.sep) + 1
            if level > 1:
                indent = "  " * (level - 1) + "+-- "
            else:
                indent = ""

            for dir_name in sorted(dirs):
                structure.append(f"{indent}{dir_name}/")

            sub_indent = "  " * level + "+-- "
            for file in sorted(files):
                if not file.startswith("."):
                    structure.append(f"{sub_indent}{file}")

        return "\n".join(structure)

    incompatible_files = []
    incompatible_structure = []
    output_lines = []

    timestamp = datetime.now().strftime("%d.%m.%Y_%H_%M_%S")
    base_filename = f"project_documentation_{timestamp}"
    output_dir = os.path.join(project_path, base_filename)
    os.makedirs(output_dir, exist_ok=True)
    output_file_with_timestamp = os.path.join(output_dir, f"{base_filename}.txt")

    # Create individual files for selected files
    if main_window.selected_files_for_compression:
        output_dir_selected = os.path.join(output_dir, "selected_files_content")
        os.makedirs(output_dir_selected, exist_ok=True)
        for file_path in main_window.selected_files_for_compression:
            try:
                content = _read_file_content(file_path)  # Call _read_file_content directly
                output_file_name = os.path.basename(file_path) + ".txt"
                output_path = os.path.join(output_dir_selected, output_file_name)
                with open(output_path, 'w', encoding='utf-8') as outfile:
                    outfile.write(content)
                print(f"Created individual file: {output_path}")
            except Exception as e:
                print(f"Error creating individual file for {file_path}: {e}")
                QMessageBox.warning(main_window, "Error", f"Error creating individual file for {os.path.basename(file_path)}.")

    output_lines.append(f"# Project Documentation: {os.path.basename(project_path)}\n")
    output_lines.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

    if llm_overview:
        output_lines.append("## PROJECT GENERAL OVERVIEW\n\n")
        output_lines.append(llm_overview + "\n\n")

    output_lines.append("## Project Structure\n\n```\n")
    structure = get_project_structure(project_path)
    output_lines.append(structure)
    output_lines.append("\n```\n\n")

    output_lines.append("## Files Content\n\n")

    for root, dirs, files in os.walk(project_path):
        print(f"Processing root: {root}")
        print(f"Current dirs: {dirs}")
        print(f"Current files: {files}")

        dirs[:] = [d for d in dirs if d != "__pycache__" and not d.startswith(".")]
        for file in files:
            if file.startswith("."):
                continue

            file_lines = main_window._process_single_file(root, file, project_path)
            output_lines.extend(file_lines)

            # Handle incompatible files
            if not file_lines or file.endswith((".wasm", ".snap")):
                rel_path = os.path.relpath(os.path.join(root, file), project_path)
                incompatible_files.append(rel_path)
                level = rel_path.count(os.sep)
                indent = "  " * level
                incompatible_structure.append(f"{indent}{os.path.basename(rel_path)}")

    if incompatible_files:
        incompatible_file_path = os.path.join(output_dir, "incompatible_files.txt")
        with open(incompatible_file_path, "w", encoding="utf-8") as incompatible_file:
            incompatible_file.write("Incompatible files:\n")
            incompatible_file.write("\n".join(incompatible_files))

        QMessageBox.information(
            main_window,
            "Info",
            "Documentation generated with incompatible files.\n"
            f"Details in: {incompatible_file_path}",
        )

    try:
        with open(output_file_with_timestamp, "w", encoding="utf-8") as outfile:
            outfile.writelines(output_lines)

        QMessageBox.information(
            main_window,
            "Success",
            "Documentation generated successfully!\n"
            f"Saved to: {output_file_with_timestamp}",
        )
        return True

    except Exception as e:
        print(f"Error during text conversion: {e}")
        QMessageBox.critical(main_window, "Error", f"An error occurred: {str(e)}")
        return False

def create_project_documentation(main_window, project_path, llm_content=None):
    """
    Creates detailed project documentation in DOCX format, handling more extensions and recognizing code blocks.
    """
    print(
        f"Creating DOCX project documentation for: {project_path}"
    )
    timestamp = datetime.now().strftime("%d.%m.%Y_%H_%M_%S")
    base_filename = f"project_documentation_{timestamp}"
    output_dir = os.path.join(project_path, base_filename)
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, f"{base_filename}.docx")
    print(f"Output file: {output_file}")
    doc = Document()

    # Define styles
    styles = doc.styles

    style_normal = styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    for level in range(1, 4):
        style_name = f"Custom Heading {level}"
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles[f"Heading {level}"]
            style.font.name = "Calibri"
            style.font.size = Pt(16 - level)
            style.font.bold = True

    doc.add_heading("Project Documentation", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Project Path: {project_path}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Document sections:", style="List Bullet")
    sections = [
        "1. Project Overview",
        "2. Project Structure",
        "3. Code and Log Files Documentation",
        "4. Incompatible Files",
        "5. Dependencies",
        "6. Setup Instructions",
    ]
    if llm_content:
        sections.insert(4, "7. LLM Analysis")
    for section in sections:
        doc.add_paragraph(section, style="List Number")
    doc.add_page_break()

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This documentation provides a comprehensive overview of the project structure and contents."
    )
    doc.add_paragraph("Project Details:", style="Custom Heading 3")
    project_name = os.path.basename(project_path)
    details = [
        f"Project Name: {project_name}",
        f'Documentation Date: {datetime.now().strftime("%Y-%m-%d")}',
        f'Number of Python Files: {sum(1 for root, _, files in os.walk(project_path) for file in files if file.endswith(".py"))}',
        f'Number of Log Files: {sum(1 for root, _, files in os.walk(project_path) for file in files if file.endswith(".log"))}',
    ]
    for detail in details:
        doc.add_paragraph(detail, style="List Bullet")

    doc.add_heading("2. Project Structure", level=1)
    doc.add_paragraph("Directory structure of the project:", style="Custom Heading 3")

    structure = []
    incompatible_files = []
    incompatible_structure = []
    for root, dirs, files in os.walk(project_path):
        dirs[:] = [d for d in dirs if d != "__pycache__" and not d.startswith(".")]
        level = root.replace(project_path, "").count(os.sep)
        indent = "    " * level
        folder = os.path.basename(root)
        structure.append(f"{indent}{folder}/")
        for file in sorted(files):
            if not file.startswith("."):
                if has_any_extension(file, [".py", ".json", ".log", ".yaml", ".md", ".ts", ".mjs", ".toml", ".txt", ".htm", ".html"]):
                    structure.append(f"{indent}    {file}")
                else:
                    incompatible_files.append(os.path.relpath(os.path.join(root, file), project_path))
                    incompatible_structure.append(f"{indent}    {file}")
    doc.add_paragraph().add_run("\n".join(structure)).font.name = "Courier New"

    # Create individual files for selected files within the documentation folder
    if main_window.selected_files_for_compression:
        output_dir_selected = os.path.join(output_dir, "selected_files_content")
        os.makedirs(output_dir_selected, exist_ok=True)
        for file_path in main_window.selected_files_for_compression:
            try:
                content = _read_file_content(file_path)  # Call _read_file_content directly
                output_file_name = os.path.basename(file_path) + ".txt"
                output_path = os.path.join(output_dir_selected, output_file_name)
                with open(output_path, 'w', encoding='utf-8') as outfile:
                    outfile.write(content)
                print(f"Created individual file: {output_path}")
            except Exception as e:
                print(f"Error creating individual file for {file_path}: {e}")
                QMessageBox.warning(main_window, "Error", f"Error creating individual file for {os.path.basename(file_path)}.")

    doc.add_heading("3. Code and Log Files Documentation", level=1)

    for root, _, files in os.walk(project_path):
        for file in files:
            if has_any_extension(file, [".py", ".json", ".log", ".yaml", ".md", ".ts", ".mjs", ".toml", ".txt", ".htm", ".html"]) and not file.startswith("."):
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, project_path)

                doc.add_heading(f"File: {rel_path}", level=2)

                try:
                    content = _read_file_content(file_path)  # Call _read_file_content directly
                    p = doc.add_paragraph()
                    run = p.add_run(content)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"Error reading file: {str(e)}")

                doc.add_paragraph()

    doc.add_heading("4. Incompatible Files", level=1)
    if incompatible_files:
        doc.add_paragraph("List of incompatible files:", style="Custom Heading 3")
        doc.add_paragraph().add_run("\n".join(incompatible_structure)).font.name = "Courier New"
        doc.add_paragraph("Incompatible files details:", style="List Bullet")
        for file_path in incompatible_files:
            doc.add_paragraph(file_path, style="List Bullet")
    else:
        doc.add_paragraph("No incompatible files found.")

    if llm_content:
        doc.add_heading("7. LLM Analysis", level=1)
        doc.add_paragraph(llm_content)
    else:
        doc.add_paragraph("LLM documentation was not requested for this document.")

    doc.add_heading("5. Dependencies", level=1)
    doc.add_paragraph("List of potential project dependencies:", style="Custom Heading 3")
    doc.add_paragraph("To be filled manually with:")
    dependencies = [
        "Required Python version",
        "Required external packages",
        "System requirements",
        "Additional software dependencies",
    ]
    for dep in dependencies:
        doc.add_paragraph(dep, style="List Bullet")

    doc.add_heading("6. Setup Instructions", level=1)
    doc.add_paragraph("Template for setup instructions:", style="Custom Heading 3")
    instructions = [
        "Environment setup",
        "Installation steps",
        "Configuration requirements",
        "Running the project",
        "Testing procedures",
    ]
    for instruction in instructions:
        doc.add_paragraph(instruction, style="List Bullet")

    try:
        doc.save(output_file)
        QMessageBox.information(
            main_window,
            "Success",
            "Documentation generated successfully!\n"
            f"Saved to: {output_file}"
        )
        print(f"DOCX documentation generated successfully at: {output_file}")
    except Exception as e:
        QMessageBox.critical(
            main_window,
            "Error",
            f"An error occurred while saving: {str(e)}"
        )
        print(f"Error during DOCX generation: {str(e)}")

def convert_docx_to_txt(docx_path):
    """
    Converts a DOCX file to a TXT file.

    Args:
        docx_path (str): Path to the DOCX file.

    Returns:
        str: Path to the converted TXT file, or None if an error occurred.
    """
    try:
        doc = Document(docx_path)
        txt_path = os.path.join(
            os.path.dirname(docx_path),
            os.path.splitext(os.path.basename(docx_path))[0] + ".txt"
        )

        with open(txt_path, "w", encoding="utf-8") as txt_file:
            for paragraph in doc.paragraphs:
                txt_file.write(paragraph.text + "\n")

        return txt_path

    except Exception as e:
        print(f"Error converting DOCX to TXT: {str(e)}")
        return None
